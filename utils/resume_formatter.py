from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ✅ Save as DOCX
def save_resume_docx(content, filename="resume.docx"):
    doc = Document()
    doc.add_heading("Resume", 0)
    doc.add_paragraph(content)
    doc.save(filename)
    return filename

# ✅ Save as PDF
def save_resume_pdf(content, filename="resume.pdf"):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    text = c.beginText(40, height - 50)
    text.setFont("Helvetica", 12)

    for line in content.split("\n"):
        text.textLine(line)
    c.drawText(text)
    c.save()
    return filename


# ---------- TEST CASE ----------
if __name__ == "__main__":
    sample_resume = """RESUME

    Skills: Python, SQL, ML, Data Visualization

    Summary: Candidate with strong skills in Python, SQL, and Machine Learning.
    """
    print("DOCX saved:", save_resume_docx(sample_resume))
    print("PDF saved:", save_resume_pdf(sample_resume))
